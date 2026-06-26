from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("Orbit_Teknik_Rapor.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
TEXT = RGBColor(35, 35, 35)
MUTED = RGBColor(95, 95, 95)
LIGHT_FILL = "F2F4F7"
BORDER = "B7C9D9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = BORDER) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_table(table, widths: list[float]) -> None:
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            if row_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                    if row_idx == 0:
                        run.bold = True
                        run.font.color.rgb = DARK_BLUE


def add_table(doc: Document, rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            table.cell(r, c).text = value
    style_table(table, widths)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = BLUE if level <= 2 else DARK_BLUE


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Orbit Teknik Raporu | BM 416 Oyun Programlama")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def build() -> None:
    doc = Document()
    setup_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ORBIT TEKNİK RAPORU")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Flutter + Flame Engine Mobil Arcade Oyunu")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    add_heading(doc, "1. Projenin Amacı")
    add_body(
        doc,
        "Orbit projesinin amacı, tek dokunuşla oynanabilen, kısa oturumlu ve refleks temelli bir mobil arcade oyunu geliştirmektir. Oyuncu merkezdeki çekirdek etrafında dönen gezegenin yörüngesini kontrol eder, asteroitlerden kaçar, coin toplar ve yüksek skor yapmaya çalışır.",
    )

    add_heading(doc, "2. Kullanılan Teknolojiler")
    add_table(
        doc,
        [
            ["Teknoloji", "Kullanım Amacı"],
            ["Flutter", "Mobil uygulama iskeleti, overlay arayüzleri ve Android APK üretimi"],
            ["Flame Engine 1.35.1", "Oyun döngüsü, component sistemi, collision ve render işlemleri"],
            ["flame_audio", "Tap, coin ve game over ses efektleri"],
            ["shared_preferences", "Coin, high score ve mağaza seçimlerini yerelde saklama"],
            ["Dart", "Oyun mantığı, component davranışları ve veri yönetimi"],
        ],
        [2.0, 4.5],
    )

    add_heading(doc, "3. Kod Mimarisi")
    add_body(
        doc,
        "Proje modüler bir yapıda geliştirildi. Ana oyun sınıfı genel oyun durumunu yönetirken, oyuncu, asteroit, coin ve güneş gibi oyun nesneleri ayrı component dosyalarında tutuldu.",
    )
    add_table(
        doc,
        [
            ["Dosya / Modül", "Görev"],
            ["main.dart", "GameWidget kurulumu ve Flutter overlay kayıtları"],
            ["orbital_gravity_game.dart", "Ana oyun döngüsü, spawn, skor, collision ve state yönetimi"],
            ["player_component.dart", "Oyuncu çizimi, trail, hitbox ve asteroid collision yönlendirmesi"],
            ["asteroid_component.dart", "Asteroit hareketi, kırmızı trail ve cleanup"],
            ["coin_component.dart", "Coin çizimi, yanıp sönme efekti, collision ve fade-out"],
            ["sun_component.dart", "Merkez çekirdek ve seçilebilir sun skin render işlemleri"],
            ["shop_data_store.dart", "SharedPreferences ile kalıcı veri saklama"],
        ],
        [2.35, 4.15],
    )

    add_heading(doc, "4. Temel Oyun Mekaniği Nasıl Yapıldı?")
    add_body(
        doc,
        "Oyuncunun konumu doğrudan fizik motoruyla değil, trigonometrik yörünge hesabıyla belirlendi. Merkez nokta güneşin konumu kabul edildi. Her frame açısal değer artırıldı ve oyuncu pozisyonu cos/sin fonksiyonlarıyla hesaplandı.",
    )
    add_bullets(
        doc,
        [
            "x = sun.x + currentRadius * cos(angle)",
            "y = sun.y + currentRadius * sin(angle)",
            "Ekrana basılı tutulduğunda targetRadius minimum değere çekildi.",
            "Parmak bırakıldığında targetRadius maksimum değere yöneldi.",
            "currentRadius, targetRadius değerine LERP ile yaklaştırıldı; böylece hareket yumuşak oldu.",
        ],
    )

    add_heading(doc, "5. Asteroit, Coin ve Collision Sistemi")
    add_body(
        doc,
        "Asteroitler ekranın dört kenarından, görünür alanın biraz dışından doğacak şekilde üretildi. Doğdukları noktadan merkez çekirdeğe doğru normalize edilmiş bir yön vektörü hesaplandı. Her update çağrısında asteroit bu yönde speed * dt kadar ilerletildi.",
    )
    add_bullets(
        doc,
        [
            "Asteroit-player çarpışması Flame CircleHitbox ile algılandı.",
            "Asteroit merkeze ulaşırsa sahneden silindi ve oyuncuya dodge skoru verildi.",
            "Coinler yörünge alanında rastgele açı ve yarıçapla doğdu.",
            "Coin-player çarpışmasında totalCoins +5 artırıldı ve kayıt işlemi yapıldı.",
            "Coinlerin kullanıcı tarafından fark edilmesi için alpha tabanlı yanıp sönme efekti eklendi.",
        ],
    )

    add_heading(doc, "6. Skor, Zorluk ve Dual Orbit")
    add_body(
        doc,
        "Skor sistemi iki parçadan oluşur: oyuncu hayatta kaldığı her saniye +1 skor kazanır, asteroit oyuncuya çarpmadan merkeze ulaşırsa +5 dodge skoru eklenir. Zorluk skor arttıkça asteroid spawn aralığını düşürür ve asteroid hızını artırır.",
    )
    add_table(
        doc,
        [
            ["Sistem", "Uygulama Detayı"],
            ["Başlangıç Dengesi", "İlk skor aralığında spawn yavaş ve asteroid hızı düşük tutuldu"],
            ["Zorluk Artışı", "Skora göre spawn interval düşürülür, asteroid speed artırılır"],
            ["Dual Orbit", "100 skorda ikinci oyuncu angle + pi konumunda oluşturulur"],
            ["Kolaylaştırıcılar", "Dual Orbit başlangıcında invincibility, spawn molası ve küçültülmüş hitbox kullanıldı"],
        ],
        [1.75, 4.75],
    )

    add_heading(doc, "7. Mağaza ve Kayıt Sistemi")
    add_body(
        doc,
        "Oyunda coin tabanlı basit bir ekonomi sistemi vardır. Oyuncu kazandığı coinlerle oyuncu rengi, arka plan teması, trail efekti, sun skin ve Energy Shield satın alabilir. Bu bilgiler SharedPreferences ile cihazda saklanır.",
    )
    add_bullets(
        doc,
        [
            "saveGameState update döngüsünde çağrılmadı; sadece coin toplama, mağaza işlemi ve game over anında çalıştırıldı.",
            "loadGameState oyun yüklenirken çağrıldı.",
            "High score, totalCoins, seçili kozmetikler ve unlock listesi kalıcı tutuldu.",
            "Energy Shield tek kullanımlık bir koruma olarak tasarlandı.",
        ],
    )

    add_heading(doc, "8. Performans İçin Alınan Önlemler")
    add_bullets(
        doc,
        [
            "Trail listeleri sabit uzunlukta tutuldu; oyuncu trail 6, asteroit trail 10 nokta ile sınırlandırıldı.",
            "Aynı anda sahnede bulunabilecek asteroid sayısı maksimum 24 ile sınırlandırıldı.",
            "Asteroit ve coin sahneden silinirken hitbox'ların collision ağacından ayrılması sağlandı.",
            "Starry Space temasındaki yıldız koordinatları render içinde üretilmedi; resize sırasında cache'lendi.",
            "Tap sesleri AudioPool ile oynatıldı ve 120 ms cooldown eklendi. Böylece hızlı dokunma spam'inin telefonlarda performans düşürmesi azaltıldı.",
        ],
    )

    add_heading(doc, "9. Test ve Çıktı")
    add_body(
        doc,
        "Proje Flutter analiz ve test komutlarıyla doğrulandı. Son aşamada Android release APK çıktısı alındı ve fiziksel telefonlarda denenmek üzere hazırlandı.",
    )
    add_table(
        doc,
        [
            ["Kontrol", "Sonuç"],
            ["flutter analyze", "Temiz geçti"],
            ["flutter test", "Widget smoke testi geçti"],
            ["Release APK", "build/app/outputs/flutter-apk/app-release.apk"],
            ["Manuel Test", "Birden fazla Android telefonda oynanış ve performans geri bildirimi alındı"],
        ],
        [2.0, 4.5],
    )

    add_heading(doc, "10. Sonuç")
    add_body(
        doc,
        "Orbit, tek dokunuşlu yörünge kontrolü üzerine kurulu, Flame component mimarisiyle geliştirilen, mağaza ve kayıt sistemi bulunan tamamlanmış bir mobil arcade prototipidir. Projede temel oyun döngüsü, collision, skor, zorluk eğrisi, shop, ses ve performans optimizasyonları birlikte uygulanmıştır.",
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
