from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("Orbit_GDD.docx")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_FILL = "F2F4F7"
BORDER = "B7C9D9"
TEXT = RGBColor(35, 35, 35)
MUTED = RGBColor(95, 95, 95)
GOLD = RGBColor(255, 183, 3)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = BORDER) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths_in: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            cell = row.cells[idx]
            cell.width = Inches(width)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)


def style_table(table, widths_in: list[float], header: bool = True) -> None:
    table.autofit = False
    set_table_width(table, widths_in)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            if row_idx == 0 and header:
                set_cell_shading(cell, LIGHT_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = DARK_BLUE
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)


def add_table(doc: Document, rows: list[list[str]], widths_in: list[float]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            table.cell(r, c).text = value
    style_table(table, widths_in)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = BLUE if level in (1, 2) else DARK_BLUE
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Orbit GDD  |  BM 416 Oyun Programlama")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def build() -> None:
    doc = Document()
    setup_styles(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GAME DESIGN DOCUMENT")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Orbit")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = TEXT

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Versiyon 1.0  •  Haziran 2026  •  Ders Projesi")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        "Ders: BM 416 Oyun Programlama  |  Platform: Android / iOS  |  Tür: Minimalist Mobil Arcade"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    doc.add_paragraph()

    add_heading(doc, "1. Oyun Özeti")
    add_body(
        doc,
        "Orbit, oyuncunun tek dokunuşla merkezdeki kozmik çekirdek etrafında dönen bir gezegenin yörünge yarıçapını kontrol ettiği minimalist bir mobil arcade oyunudur. Oyuncu, ekrana basılı tuttuğunda yörüngeyi daraltır; parmağını kaldırdığında yörünge genişler. Amaç, merkeze doğru akan asteroitlerden kaçmak, coin toplamak ve mümkün olduğunca yüksek skor elde etmektir.",
    )
    add_body(
        doc,
        "Oyun kısa oturumlara, refleks temelli karar vermeye ve tek parmakla hızlı oynanabilirliğe odaklanır. 100 skordan sonra Dual Orbit fazı aktif olur ve oyuncu aynı anda birbirine zıt konumda dönen iki gezegeni yöneterek daha ileri bir beceri sınavına girer.",
    )

    add_heading(doc, "2. Temel Bilgiler")
    add_table(
        doc,
        [
            ["Alan", "Detay"],
            ["Oyun Adı", "Orbit"],
            ["Tür", "Minimalist Mobil Arcade / Hyper-casual Refleks"],
            ["Platform", "Android öncelikli, iOS uyumlu yapı"],
            ["Hedef Kitle", "Kısa oturumlu refleks oyunlarını seven mobil oyuncular"],
            ["Motor", "Flutter + Flame Engine 1.35.1"],
            ["Kontrol Şeması", "Tek parmak basılı tutma / bırakma"],
            ["Oturum Süresi", "30 saniye - 3 dakika"],
            ["Görsel Stil", "Koyu uzay arka planı, neon renkler, sade geometrik şekiller"],
            ["Yayın Modeli", "Ücretsiz prototip / ders projesi"],
        ],
        [1.9, 4.6],
    )

    add_heading(doc, "3. Hikaye & Evren")
    add_heading(doc, "3.1 Genel Konsept", 2)
    add_body(
        doc,
        "Orbit soyut bir uzay evreninde geçer. Merkezdeki kozmik çekirdek sürekli olarak oyuncunun gezegenini kendine çeker. Oyuncu, bu çekimi tek dokunuşla dengeleyerek güvenli bir yörüngede kalmaya çalışır. Ekranın dışından gelen asteroitler çekirdeğe doğru düşer ve oyuncu bu tehlikeli akışın içinde hayatta kalır.",
    )
    add_heading(doc, "3.2 Tematik Unsurlar", 2)
    add_bullets(
        doc,
        [
            "Kozmik çekirdek: Oyunun merkezi ve yörünge referans noktasıdır.",
            "Oyuncu gezegeni: Tek dokunuşla kontrol edilen ana oyun nesnesidir.",
            "Asteroitler: Kırmızı renk ve iz efektiyle tehlikeyi temsil eder.",
            "Coinler: Sarı renkte, yanıp sönen güvenli ödül nesneleridir.",
            "Dual Orbit: Yüksek skor sonrası oyuncunun ustalığını test eden gelişmiş fazdır.",
        ],
    )

    add_heading(doc, "4. Oyun Mekaniği")
    add_heading(doc, "4.1 Temel Kontroller", 2)
    add_table(
        doc,
        [
            ["Aksiyon", "Mobil Girdi", "Oyun Etkisi"],
            ["Basılı Tutma", "Ekrana dokun ve tut", "Yörünge yarıçapı minimum değere doğru daralır"],
            ["Bırakma", "Parmağı ekrandan kaldır", "Yörünge yarıçapı maksimum değere doğru genişler"],
            ["Dokunmama", "Pasif durum", "Güneş çekimi oyuncuyu yavaşça merkeze doğru çeker"],
            ["Restart", "Game Over ekranı", "Aktif asteroit/coinler temizlenir, skor sıfırlanır"],
            ["Shop", "Main Menu / Game Over", "Kozmetik ve beceri satın alma ekranı açılır"],
        ],
        [1.55, 1.8, 3.15],
    )

    add_heading(doc, "4.2 Yörünge Matematiği", 2)
    add_body(
        doc,
        "Oyuncu konumu merkez koordinatları, mevcut yarıçap ve açı değeri kullanılarak trigonometrik olarak hesaplanır:",
    )
    add_bullets(
        doc,
        [
            "x = sun.x + currentRadius * cos(angle)",
            "y = sun.y + currentRadius * sin(angle)",
            "currentRadius, targetRadius değerine LERP ile yumuşak geçiş yapar.",
            "Dual Orbit aktifken ikinci oyuncu angle + pi konumunda yer alır.",
        ],
    )

    add_heading(doc, "4.3 Ana Sistemler", 2)
    add_bullets(
        doc,
        [
            "Asteroit sistemi: Asteroitler ekran kenarlarının dışından doğar ve merkeze normalize edilmiş vektörle ilerler.",
            "Coin sistemi: Coinler yörünge alanında doğar, toplanınca totalCoins değerini +5 artırır.",
            "Skor sistemi: Her saniye +1 skor, merkeze güvenli ulaşan asteroit başına +5 dodge skoru.",
            "Shield sistemi: Energy Shield satın alındıysa ilk asteroid çarpışmasını engeller.",
            "Game Over sistemi: Asteroide çarpma veya güneş sınırına düşme oyunu bitirir.",
        ],
    )

    add_heading(doc, "5. İlerleme ve Zorluk Dengesi")
    add_table(
        doc,
        [
            ["Faz", "Skor Aralığı", "Spawn / Hız Mantığı", "Tasarım Amacı"],
            ["Very Easy", "0 - 14", "3.0 sn spawn, 60 px/sn asteroid", "Oyuncunun kontrolü öğrenmesi"],
            ["Easy", "15 - 40", "2.0 sn spawn, 80 px/sn asteroid", "Düzenli refleks baskısı oluşturmak"],
            ["Hard Scaling", "40+", "Spawn 1.8 sn'den 0.9 sn'ye iner, hız 180 px/sn'ye kadar çıkar", "Uzun oturumlarda meydan okumayı artırmak"],
            ["Dual Orbit", "100+", "Zorluk score 0 seviyesine resetlenir, sonra iki kat yavaş artar", "İkinci gezegenin bilişsel yükünü dengelemek"],
        ],
        [1.25, 1.15, 2.55, 1.55],
    )
    add_body(
        doc,
        "Dual Orbit etkinleştiğinde oyuncuya 2 saniyelik asteroid spawn molası ve 3 saniyelik çarpışma bağışıklığı verilir. Ayrıca iki oyuncunun hitbox yarıçapı %45 küçültülerek görsel olarak yakın geçen asteroitlerin haksız game over üretmesi engellenir.",
    )

    add_heading(doc, "6. Seviye Tasarımı")
    add_body(
        doc,
        "Orbit klasik bölüm yapısı yerine sonsuz arena yapısı kullanır. Tüm oynanış tek ekranda, merkezdeki çekirdek etrafında gerçekleşir. Seviye tasarımının ana değişkenleri asteroid spawn yönü, asteroid hızı, spawn aralığı, coin konumu ve oyuncunun yörünge yarıçapıdır.",
    )
    add_table(
        doc,
        [
            ["Alan", "Açıklama"],
            ["Oyun Alanı", "Dikey mobil ekran, merkezde sabit kozmik çekirdek"],
            ["Spawn Kenarları", "Üst, alt, sol ve sağ ekran dışı alanlar"],
            ["Coin Alanı", "60 - 200 px yörünge yarıçapı içinde rastgele açı"],
            ["Tema Varyasyonları", "Deep Cosmic Black, Starry Space, Cyberpunk Grid"],
            ["Milestone", "100 skor sonrası Dual Orbit fazı"],
        ],
        [1.7, 4.8],
    )

    add_heading(doc, "7. Nesne ve Düşman Tasarımı")
    add_table(
        doc,
        [
            ["Nesne", "Tür", "Davranış", "Oyuncu Etkisi"],
            ["Player", "Kontrol edilen gezegen", "Merkez etrafında döner, yarıçap dokunuşla değişir", "Hayatta kalma ana hedefidir"],
            ["Player 2", "Dual Orbit gezegeni", "100 skor sonrası zıt açıda doğar", "Zorluğu ve stratejik takip ihtiyacını artırır"],
            ["Sun / Core", "Merkez nesne", "Çekim tehdidi ve görsel odak noktasıdır", "Yarıçap çok düşerse game over"],
            ["Asteroid", "Düşman", "Ekran dışından merkeze doğru akar, kırmızı trail bırakır", "Çarpışma game over veya shield tüketimi"],
            ["Coin", "Ödül", "Yörünge alanında doğar, sarı yanıp söner", "+5 totalCoins"],
            ["Energy Shield", "Skill", "Shop'tan satın alınan tek kullanımlık koruma", "İlk asteroid çarpışmasını engeller"],
        ],
        [1.2, 1.25, 2.45, 1.6],
    )

    add_heading(doc, "8. Sanat Yönetimi & Ses")
    add_heading(doc, "8.1 Görsel Stil", 2)
    add_bullets(
        doc,
        [
            "Ana stil: Minimalist neon uzay estetiği.",
            "Arka plan: Deep space siyahı (#0B0C10).",
            "Oyuncu: Varsayılan neon cyan (#45A29E), shop ile yeşil veya mor seçenekler.",
            "Asteroit: Neon kırmızı/turuncu (#FF4D3D) ve belirgin kırmızı trail.",
            "Coin: Sarı/altın (#FFD700), kullanıcıya ödül olduğunu göstermek için yanıp söner.",
            "Sun skin seçenekleri: Standard Core, Black Hole Core, Solar Eclipse.",
        ],
    )
    add_heading(doc, "8.2 Müzik & Ses", 2)
    add_bullets(
        doc,
        [
            "Tap sesi: Oyuncu dokunduğunda kısa ve düşük hacimli feedback.",
            "Coin sesi: Coin toplandığında kısa ödül sesi.",
            "Explosion sesi: Game over anında düşük ve etkili patlama/glitch sesi.",
            "Performans için sesler AudioPool üzerinden oynatılır; tap sesi 120 ms cooldown ile sınırlandırılır.",
        ],
    )

    add_heading(doc, "9. Arayüz (UI/UX)")
    add_table(
        doc,
        [
            ["UI Elementi", "Konum", "Açıklama"],
            ["Main Menu", "Overlay", "ORBIT başlığı, START GAME ve SHOP butonları"],
            ["Score Text", "Üst orta", "Oyun sırasında canlı skor gösterimi"],
            ["Game Over", "Overlay", "Final score, high score, Try Again ve Shop butonları"],
            ["Advanced Shop", "Overlay", "Kategori kartları ve item grid düzeni"],
            ["Dual Orbit Uyarısı", "Oyun alanı üst bölge", "Dual Orbit başladığında kısa bilgilendirme metni"],
            ["Coin Balance", "Shop üst bölge", "Oyuncunun toplam coin miktarını gösterir"],
        ],
        [1.55, 1.35, 3.6],
    )

    add_heading(doc, "10. Ekonomi & Mağaza")
    add_body(
        doc,
        "Oyun içi ekonomi yalnızca kozmetik ve tek kullanımlık beceri satın alımlarına hizmet eder. Coinler oynanış sırasında toplanır ve SharedPreferences ile kalıcı olarak saklanır.",
    )
    add_table(
        doc,
        [
            ["Kategori", "Item", "Fiyat", "Etkisi"],
            ["Orbit Colors", "Neon Cyan", "0", "Varsayılan oyuncu rengi"],
            ["Orbit Colors", "Neon Green", "50", "Oyuncu rengini yeşil yapar"],
            ["Orbit Colors", "Neon Purple", "100", "Oyuncu rengini mor yapar"],
            ["Backgrounds", "Starry Space", "100", "Arka plana küçük statik yıldızlar ekler"],
            ["Backgrounds", "Cyberpunk Grid", "150", "Koyu grid arka planı"],
            ["Skills", "Energy Shield", "200", "Bir sonraki run için tek kullanımlık koruma"],
            ["Trails", "Rainbow Trail", "150", "Oyuncu trail rengini dinamik yapar"],
            ["Trails", "Star Dust Trail", "250", "Trail'i küçük yıldız parçacıkları gibi çizer"],
            ["Cosmic Cores", "Black Hole Core", "300", "Güneşi vortex/black hole stiline çevirir"],
            ["Cosmic Cores", "Solar Eclipse", "400", "Güneşi eclipse temalı çekirdeğe çevirir"],
        ],
        [1.45, 1.75, 0.75, 2.55],
    )

    add_heading(doc, "11. Teknik Gereksinimler")
    add_table(
        doc,
        [
            ["Özellik", "Minimum", "Önerilen"],
            ["İşletim Sistemi", "Android 8.0+", "Android 10+"],
            ["Cihaz", "Orta seviye telefon", "Güncel orta/üst seviye Android telefon"],
            ["RAM", "3 GB", "4 GB+"],
            ["FPS Hedefi", "Stabil 45+ FPS", "60 FPS"],
            ["Motor", "Flutter + Flame 1.35.1", "Release APK ile test"],
            ["Depolama", "100 MB altı APK", "Yeterli boş alan ve güncel Play Services"],
        ],
        [1.55, 2.35, 2.6],
    )
    add_bullets(
        doc,
        [
            "Trail listeleri sabit uzunlukla sınırlandırılmıştır.",
            "Asteroit sayısı aynı anda maksimum 24 olacak şekilde sınırlandırılmıştır.",
            "Starry Space yıldız koordinatları render sırasında random üretilmez; resize/init sırasında cache'lenir.",
            "SharedPreferences kayıtları update loop içinde değil, yalnızca coin toplama, shop işlemi ve game over anında yapılır.",
        ],
    )

    add_heading(doc, "12. Geliştirme Takvimi")
    add_table(
        doc,
        [
            ["Aşama", "Dönem", "Hedefler"],
            ["Phase 1", "Tamamlandı", "Yörünge matematiği, sun, player, touch input, LERP"],
            ["Phase 2", "Tamamlandı", "Asteroit spawn sistemi, ekran dışından merkeze hareket"],
            ["Phase 3", "Tamamlandı", "Collision, skor, health/game over, denge eğrisi"],
            ["Phase 4", "Tamamlandı", "UI overlay, shop, coin, persistence, görsel polish"],
            ["Hardening", "Tamamlandı", "AudioPool, performans limitleri, APK release testleri"],
            ["Teslim", "Haziran 2026", "GDD, APK ve proje dosyalarının ders teslimine hazırlanması"],
        ],
        [1.35, 1.55, 3.6],
    )

    add_heading(doc, "13. Ekip")
    add_table(
        doc,
        [
            ["Rol", "Kişi / Açıklama"],
            ["Oyun Tasarımı", "[Ad Soyad]"],
            ["Programlama", "[Ad Soyad]"],
            ["UI/UX Tasarımı", "[Ad Soyad]"],
            ["Ses ve Görsel Düzenleme", "[Ad Soyad]"],
            ["Test", "Farklı Android telefonlarda manuel test"],
            ["Ders", "BM 416 Oyun Programlama"],
        ],
        [2.1, 4.4],
    )

    add_heading(doc, "14. Riskler ve Çözüm Notları")
    add_bullets(
        doc,
        [
            "Telefonlarda ses spam kaynaklı kasma riskine karşı AudioPool ve tap cooldown eklendi.",
            "Uzun oynanışta obje birikmesini önlemek için asteroid ve coin cleanup akışları güçlendirildi.",
            "Dual Orbit fazı zor olabileceği için zorluk reseti, 3 saniye invincibility ve küçültülmüş hitbox kullanıldı.",
            "Shop seçimleri ve coin bakiyesi uygulama kapanıp açılsa da kalıcı tutulur.",
        ],
    )

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = closing.add_run("Bu doküman BM 416 Oyun Programlama dersi proje teslimi için hazırlanmıştır.")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
